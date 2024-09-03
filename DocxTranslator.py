import docx
import fasttext
import gc
import os
import torch
from collections import Counter
from pathlib import Path
from transformers import pipeline
from nltk.tokenize import sent_tokenize
from zipfile import BadZipFile
import nltk
import json
from TranslatorInterface import TranslatorInterface, SupportedFileType
from datasets import Dataset, DatasetDict


class DocxTranslator(TranslatorInterface):

    def __init__(self, pretrained_lang_model="./lid218e.bin", ):
        nltk.download('punkt')
        nltk.download('punkt_tab')

        self.ext_in = SupportedFileType.DOCX.value
        self.ext_out = SupportedFileType.DOCX.value
        self.pretrained_lang_model = pretrained_lang_model

        with open("./constants/model_checkpoints.json", 'r') as json_file:
            self.checkpoints = json.load(json_file)

        with open("./constants/languages.json", 'r') as json_file:
            self.languages = json.load(json_file)

        self._prepare_language_detection_model()


    def translate(self, filePath, source_language, target_language, model_name):
        print("translating")
        file = filePath.resolve()

        if target_language not in self.languages:
            print(f"Cannot translate. {target_language} is not a supported target language.")
            return
        
        if model_name not in self.checkpoints:
            print(f"Cannot translate. {model_name} is not a supported model.")
            return 

        try:
            document = docx.Document(file)
        except BadZipFile:
            print(f"BadZipFile Error on opening {file}")
            return
        
        all_text = self.accumulate_text(file)

        languages_in_file = self._get_languages(all_text)
        top_language_in_file = languages_in_file.most_common(1)[0][0]
        file_is_src_lang = (top_language_in_file == source_language)

        if not file_is_src_lang:
            print(f"Cannot translate. File is in {top_language_in_file} expected {source_language}")
            return

        print(f"Loading model: {self.checkpoints[model_name]}")
        model, tokenizer = self._load_model(self.checkpoints[model_name])

        file_name = self.languages[target_language]
        output_dir_for_model = Path(f'./Translated/{model_name}')
        output_dir_for_model.mkdir(parents=True, exist_ok=True)
        output_path = output_dir_for_model / f"{file.stem}_{file_name}.{self.ext_out}"

        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        translation_pipeline = pipeline('translation',
                                        model=model,
                                        tokenizer=tokenizer,
                                        src_lang=source_language,
                                        tgt_lang=target_language,
                                        max_length=400,
                                        device=device)

        # Create a dataset from the text
        dataset = Dataset.from_dict({"text": all_text})

        def translate_batch(batch):
            translations = translation_pipeline(batch['text'], batch_size=16)
            return {'translated_text': [t['translation_text'] for t in translations]}

        # Translate the texts in batches
        translated_dataset = dataset.map(translate_batch, batched=True, batch_size=16)

        translation_map = dict(zip(all_text, translated_dataset['translated_text']))
        self._update_document_from_translation_map(translation_map, document)

        # Save the translated document
        document.save(output_path)

        # Frees loaded model and tokenizer
        self._unload_model(model, tokenizer)

        # Frees the pipeline
        del translation_pipeline
        gc.collect()
        torch.cuda.empty_cache()

        return output_path


    def accumulate_text(self, file):
        """
        Accumulates all text from paragraphs and tables in the DOCX file into an array of strings.
        
        Args:
            file (Path): Path to the DOCX file.
        
        Returns:
            List[str]: A list of strings where each string contains text from either a paragraph or a cell in a table.
        """

        document = docx.Document(file)
        text_accumulator = []

        # Accumulate text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():  # Only accumulate if the paragraph is not empty
                text_accumulator.append(paragraph.text)

        # Accumulate text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():  # Only accumulate if the paragraph is not empty
                            text_accumulator.append(paragraph.text)

        return text_accumulator


    def _get_languages(self, all_text):
        return self._detect_language_with_dataset(all_text)


    def _detect_language_with_dataset(self, texts):
        """
        Detects the language of each text segment using a dataset for parallel processing.
        
        Args:
            texts (List[str]): List of text segments to check language.
        
        Returns:
            Counter: A Counter object with the detected languages.
        """
        fasttext_model = fasttext.load_model(self.pretrained_lang_model)

        # Define a function to detect the language for a batch of texts
        def detect_language(batch):
            cleaned_texts = [text.replace('\n', ' ') for text in batch['text']]  # Replace newlines with spaces
            predictions = [fasttext_model.predict(text, k=1)[0][0].replace('__label__', '') for text in cleaned_texts]
            return {'language': predictions}

        # Create a dataset from the accumulated texts
        dataset = Dataset.from_dict({"text": texts})

        # Apply the language detection function to the dataset in parallel
        language_dataset = dataset.map(detect_language, batched=True, batch_size=32)

        # Accumulate the detected languages
        language_counter = Counter(language_dataset['language'])

        del fasttext_model
        gc.collect()
        torch.cuda.empty_cache()

        return language_counter


    def _replace_text_in_runs(self, paragraph, translated_text):
        """Replace text in each run while preserving the original formatting."""
        original_text = "".join(run.text for run in paragraph.runs)
        # Ensure we correctly replace text while preserving formatting
        if len(original_text) == len(translated_text):
            current_char_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = translated_text[current_char_index:current_char_index + run_length]
                current_char_index += run_length
        else:
            # If lengths don't match, replace text by matching run lengths
            current_char_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = translated_text[current_char_index:current_char_index + run_length]
                current_char_index += run_length

            # Handle any leftover text by adding it as a new run
            if current_char_index < len(translated_text):
                remaining_text = translated_text[current_char_index:]
                paragraph.add_run(remaining_text)

    
    def _update_document_from_translation_map(self, translation_map, document):
        # Replace text in the document
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                if original_text in translation_map:
                    self._replace_text_in_runs(paragraph, translation_map[original_text])

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            original_text = paragraph.text
                            if original_text in translation_map:
                                self._replace_text_in_runs(paragraph, translation_map[original_text])


# # Example usage:
translator = DocxTranslator()
file = Path("./Input/test2.docx")
translator.translate(file, "eng_Latn", "spa_Latn", "NLLB-distilled")
