import docx
import fasttext
import gc
import os
import torch
from collections import Counter
from pathlib import Path
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline, TRANSFORMERS_CACHE
from nltk.tokenize import sent_tokenize
from zipfile import BadZipFile
import nltk
import json

nltk.download('punkt')
nltk.download('punkt_tab')

class Translator:

    def __init__(self, input_folder, output_folder, ext_in='docx', ext_out='docx', pretrained_lang_model="./lid218e.bin"):
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.ext_in = ext_in
        self.ext_out = ext_out
        self.pretrained_lang_model = pretrained_lang_model

        with open("./constants/model_checkpoints.json", 'r') as json_file:
            self.checkpoints = json.load(json_file)
            # "NLLB": "facebook/nllb-200-3.3B",
            # "MADLAD": "google/madlad400-3b-mt",
            # "Llama-3.1-405B": "meta-llama/Meta-Llama-3.1-405B"

        with open("./constants/languages.json", 'r') as json_file:
            self.languages = json.load(json_file)
                # "Ethiopian": "amh_Ethi",
                # "Arabic": "arb_Arab",
                # "Assamese": "asm_Beng",
                # "Bangal": "ben_Beng",
                # "BPortugese": "por_Latn",
                # "Burmese": "mya_Mymr",
                # "Cebuano": "ceb_Latn",
                # "Chinese": "zsm_Latn",
                # "French": "fra_Latn",
                # "Gujarati": "guj_Gujr",
                # "Hausa": "hau_Latn",
                # "Hindi": "hin_Deva",
                # "Illocano": "ilo_Latn",
                # "Indonesian": "ind_Latn",
                # "Kannada": "kan_Knda",
                # "Khmer": "khm_Khmr",
                # "Laotian": "lao_Laoo",
                # "LASpanish": "spa_Latn",
                # "Malayalam": "mal_Mlym",
                # "Nepali": "npi_Deva",
                # "Oriya": "ory_Orya",
                # "PlatMalagasy": "plt_Latn",
                # "EPunjabi": "pan_Guru",
                # "Russian": "rus_Cyrl",
                # "Swahili": "swh_Latn",
                # "Tagalog": "tgl_Latn",
                # "Tamil": "tam_Taml",
                # "Telugu": "tel_Telu",
                # "Thai": "tha_Thai",
                # "TokPisin": "tpi_Latn",
                # "Urdu": "urd_Arab",

        self._prepare_language_detection_model()


    def _prepare_language_detection_model(self):
        if not os.path.isfile(self.pretrained_lang_model):
            # Download the model if it doesn't exist
            os.system(f"wget https://dl.fbaipublicfiles.com/nllb/lid/lid218e.bin")


    def load_model(self, model_name):
        model_dir = f"{TRANSFORMERS_CACHE}/{model_name}"

        if not os.path.exists(model_dir):
            print(f"{model_name} not found")
            tokenizer = AutoTokenizer.from_pretrained(model_name)
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
            model.save_pretrained(model_dir)
            tokenizer.save_pretrained(model_dir)
        else:
            print(f"{model_name} found")
            tokenizer = AutoTokenizer.from_pretrained(f"{model_dir}")
            model = AutoModelForSeq2SeqLM.from_pretrained(f"{model_dir}")

        return model, tokenizer


    def unload_model(self, model, tokenizer):
        del model
        del tokenizer
        gc.collect()
        torch.cuda.empty_cache()


    def translate_text(self, text, src_lang, tgt_lang, model_name="NLLB"):
        """Translate a given text from src_lang to tgt_lang using the specified model."""
        model, tokenizer = self.load_model(self.checkpoints[model_name])
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

        translation_pipeline = pipeline('translation',
                                        model=model,
                                        tokenizer=tokenizer,
                                        src_lang=src_lang,
                                        tgt_lang=tgt_lang,
                                        max_length=400,
                                        device=device)

        translated_text = translation_pipeline(text)[0]['translation_text']

        self.unload_model(model, tokenizer)
        return translated_text


    def translate_docx(self, input_file, output_file, model_name="NLLB"):
        """Translate a DOCX file and save the translated content to a new file."""
        model, tokenizer = self.load_model(self.checkpoints[model_name])
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

        translation_pipeline = pipeline('translation',
                                        model=model,
                                        tokenizer=tokenizer,
                                        src_lang='eng_Latn',
                                        tgt_lang='fra_Latn',
                                        max_length=400,
                                        device=device)

        doc = docx.Document(input_file)

        # Iterate over paragraphs and tables to translate the content
        for paragraph in doc.paragraphs:
            self.translate_paragraph(paragraph, translation_pipeline)

        for table in doc.tables:
            self.translate_table(table, translation_pipeline)

        # Save the translated document
        doc.save(output_file)
        self.unload_model(model, tokenizer)


    def translate_paragraph(self, paragraph, translation_pipeline):
        """Translate the content of a paragraph and replace its text."""
        original_text = paragraph.text
        if original_text.strip():  # Only translate if the paragraph is not empty
            translated_text = translation_pipeline(original_text)[0]['translation_text']
            self.replace_text_in_runs(paragraph, translated_text)


    def replace_text_in_runs(self, paragraph, translated_text):
        """Replace text in each run while preserving the original formatting."""
        original_text = "".join(run.text for run in paragraph.runs)

        # Ensure we correctly replace text while preserving formatting
        if len(original_text) == len(translated_text):
            current_char_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = translated_text[current_char_index:current_char_index + run_length]
                current_char_index += run_length
        else:
            # If lengths don't match, replace text by matching run lengths
            current_char_index = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run.text = translated_text[current_char_index:current_char_index + run_length]
                current_char_index += run_length

            # Handle any leftover text by adding it as a new run
            if current_char_index < len(translated_text):
                remaining_text = translated_text[current_char_index:]
                paragraph.add_run(remaining_text)


    def translate_table(self, table, translation_pipeline):
        """Translate all the cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.translate_paragraph(paragraph, translation_pipeline)


    def get_languages(self, file):
        file = file.resolve()
        fasttext_model = fasttext.load_model(self.pretrained_lang_model)
        
        try :
            document = docx.Document(file)
        except BadZipFile:
            print(f"BadZipFile Error on opening {file}")

        languageCounter = Counter()

        self.get_languages_in_paragraphs(document.paragraphs, fasttext_model, languageCounter)
        self.get_languages_in_tables(document.tables, fasttext_model, languageCounter)

        del fasttext_model
        gc.collect()
        torch.cuda.empty_cache()

        return languageCounter


    def get_languages_in_paragraphs(self, paragraphs, model, counter):
        sentences = [sentence for para in paragraphs for sentence in sent_tokenize(para.text)]

        for sentence in sentences:
            predictions = model.predict(sentence, k=1)
            output_lang = predictions[0][0].replace('__label__', '')
            counter.update([output_lang])


    def get_languages_in_tables(self, tables, model, counter):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self.get_languages_in_paragraphs(cell.paragraphs, model, counter)


    def translate_files(self):
        files = [file for file in self.input_folder.rglob("*." + self.ext_in)]
        print(f"Found {len(files)} {self.ext_in} files in {self.input_folder.resolve()}")

        for i, file in enumerate(files, 1):
            file = file.resolve()
            languages_in_file = self.get_languages(file)
            top_language_in_file = languages_in_file.most_common(1)[0][0]
            file_is_english = top_language_in_file == "eng_Latn"

            if file_is_english:
                print(f"{i:>4} : Translating file {file} from English to multiple languages.")
                try:
                    document = docx.Document(file)
                except BadZipFile:
                    print(f"BadZipFile Error on opening {file}")
                    continue

                for model_name, checkpoint in self.checkpoints.items():
                    print(f"Loading model: {model_name}")
                    model, tokenizer = self.load_model(checkpoint)

                    for target_lang, file_name in self.languages.items():
                        output_dir_for_model = self.output_folder / f"{model_name}"
                        output_dir_for_model.mkdir(parents=True, exist_ok=True)
                        output_path = output_dir_for_model / f"{file.stem}_{file_name}.{self.ext_out}"

                        self.translate_docx(file, output_path, model_name)

                        print(f"{i:>4} : Translated file {file} to {file_name}.")

                    self.unload_model(model, tokenizer)

            else:
                print(f"{i:>4} : Not translating file {file}. It seems to be in :{top_language_in_file}.")


# Example usage:
# translator = Translator(input_folder="./Input", output_folder="./Translated/")
# Example of translating text
# translated_text = translator.translate_text("Hello, world!", "eng_Latn", "fra_Latn", "NLLB-distilled")
# print(translated_text)

# # Translate files in a folder
# translator.translate_files()
